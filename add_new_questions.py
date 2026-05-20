"""Parse new question sources and merge into questions_data.json.

Run once from the project root:
    python add_new_questions.py [--dry-run]

Reads 4 docx files, deduplicates against existing 161 questions (SequenceMatcher > 0.82),
then appends new questions into questions_data.json starting from ID 162.
"""

import json
import os
import re
import sys
import difflib
import argparse
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

BASE_DIR = r"C:\Users\15348\Desktop\比赛\创新杯\地震测井组题库整理\地震测井组题库整理"
OUT_JSON = os.path.join(os.path.dirname(os.path.abspath(__file__)), "questions_data.json")

SEISMIC_KW = [
    "地震", "反射", "折射", "时距", "叠加", "偏移", "道集", "震源", "检波器",
    "面波", "滤波", "子波", "反褶积", "动校正", "静校正", "共炮点", "共反射点",
    "菲涅尔", "绕射", "折射波", "纵波", "横波", "波速", "波阻抗", "频散",
    "炮集", "覆盖次数", "叠前", "叠后", "同相轴", "时间剖面", "大地滤波",
    "视速度", "时差", "道间距",
]
WELL_KW = [
    "测井", "电阻率", "孔隙度", "渗透率", "饱和度", "自然电位", "伽马",
    "声波测井", "密度测井", "泥质", "侵入", "微电极", "感应测井", "泥浆",
    "钻井", "地层水", "井径", "核磁", "补偿中子", "泥岩基线", "油层",
    "储层", "岩心", "气层", "侧向测井",
]

AI_TAG = "⚠️AI生成，置信度90%"


# ─── helpers ──────────────────────────────────────────────────────────────────

def detect_category(text: str) -> str:
    s = sum(1 for kw in SEISMIC_KW if kw in text)
    w = sum(1 for kw in WELL_KW if kw in text)
    return "测井" if w > s else "地震"


def is_option_line(text: str) -> bool:
    """True if line starts with an ABCD option marker."""
    # Standard: A. / A、/ A） / (A) etc.
    if re.match(r"^[（(]?[A-D][)）．\.\、:：,，]\s*\S", text):
        return True
    # Bare letter + space: "A 界面下部介质..."
    if re.match(r"^[A-D]\s+\S", text):
        return True
    return False


def parse_opts_from_single_line(line: str) -> list[str]:
    """Parse options in formats:
    - 'A．opt  B．opt  C．opt'  (tab or space separated, standard)
    - '(A)opt  (B)opt'           (parenthesized)
    - 'A opt   B opt'            (bare letter + space, no punctuation)
    """
    # Split on option markers — look-behind avoids splitting mid-word
    parts = re.split(r"(?<!\w)(?=[（(]?[A-D][)）．\.\、:：])", line.strip())
    opts = []
    for part in parts:
        part = part.strip()
        m = re.match(r"^[（(]?([A-D])[)）．\.\、:：,，]\s*(.+)", part)
        if m:
            opts.append(f"{m.group(1)}. {m.group(2).strip().rstrip(',，')}")
    if len(opts) >= 2:
        return opts
    # Bare letter + space format: "A 界面... B 界面..."
    if re.match(r"^[A-D]\s+\S", line.strip()):
        m2 = re.match(r"^([A-D])\s+(.+)", line.strip())
        if m2:
            return [f"{m2.group(1)}. {m2.group(2).strip()}"]
    return []


def parse_opts_multiline(paras: list[str], start: int) -> tuple[list[str], int]:
    """Collect option lines starting at paras[start].
    Each line may carry 1-4 options. Returns (opts, next_index)."""
    opts: list[str] = []
    i = start
    while i < len(paras):
        text = paras[i].strip()
        if not text:
            i += 1
            continue
        if not is_option_line(text):
            break
        line_opts = parse_opts_from_single_line(text)
        if line_opts:
            opts.extend(line_opts)
        else:
            m = re.match(r"^[（(]?([A-D])[)）．\.\、:：,，]\s*(.+)", text)
            if m:
                opts.append(f"{m.group(1)}. {m.group(2).strip()}")
        i += 1
        if len(opts) >= 4:
            break
    return opts, i


def extract_answer_letters(text: str) -> str | None:
    """Find answer letter(s) embedded in text: (A), (BD), ：C, : BC."""
    # （BC） or (ABC) anywhere
    m = re.search(r"[（(]\s*([A-D]{1,4})\s*[）)]", text)
    if m:
        letters = m.group(1).upper()
        if all(c in "ABCD" for c in letters):
            return letters
    # After colon at end
    m = re.search(r"[：:]\s*([A-D]{1,4})\s*$", text)
    if m:
        return m.group(1).upper()
    # Multi-letter at end (with or without preceding whitespace)
    m = re.search(r"[\s（(]([A-D]{2,4})\s*$", text)
    if m:
        return m.group(1).upper()
    # Multi-letter immediately after Chinese char (no space): "是CD" "的BD"
    m = re.search(r"[一-鿿]([A-D]{2,4})\s*$", text)
    if m:
        return m.group(1).upper()
    # Single letter immediately after Chinese char: "是D" "为A"
    m = re.search(r"[一-鿿]\s*([A-D])\s*$", text)
    if m:
        return m.group(1).upper()
    # Single letter surrounded by wide spaces in middle: "    A    "
    m = re.search(r"\s{2,}([A-D])\s{2,}", text)
    if m:
        return m.group(1).upper()
    return None


def clean_q_text(text: str) -> str:
    """Strip answer letters/parens (anywhere) and leading numbering."""
    # Replace （letter） mid-sentence with a blank placeholder
    text = re.sub(r"\s*[（(]\s*[A-D]{1,4}\s*[）)]\s*", "____", text)
    # Strip trailing letter answers (with or without preceding whitespace/punctuation)
    text = re.sub(r"\s*[：:]\s*[A-D]{1,4}\s*$", "", text)
    text = re.sub(r"[\s。，,、]\s*[A-D]{1,4}\s*$", "", text)
    text = re.sub(r"[一-鿿]\s*[A-D]{1,4}\s*$", lambda m: m.group(0)[0], text)
    # Wide-space-surrounded letter in middle ("    A    ") → remove
    text = re.sub(r"\s{2,}[A-D]\s{2,}", " ", text)
    # Leading numbering
    text = re.sub(r"^\d{1,3}[\.、，,。]\s*", "", text)
    return text.strip().rstrip("：:").strip()


def extract_fill_answer(text: str) -> tuple[str, str]:
    """Pull answer(s) from （answer） brackets; return (answer, cleaned_q_text)."""
    answers = re.findall(r"[（(]([^（）()\n]{2,80})[）)]", text)
    if not answers:
        return "", text.strip()
    combined = "；".join(a.strip() for a in answers)
    cleaned = re.sub(r"\s*[（(][^（）()\n]{2,80}[）)]\s*", "____", text)
    cleaned = re.sub(r"^\d{1,3}[\.、，,。]\s*", "", cleaned).strip()
    return combined, cleaned


def infer_type(opts: list, answer: str) -> str:
    if not opts:
        return "fill"
    if len(answer) > 1 and all(c in "ABCD" for c in answer):
        return "multi"
    return "choice"


def is_duplicate(text: str, existing_texts: list[str], threshold: float = 0.82) -> bool:
    text = text.strip()
    for t in existing_texts:
        if difflib.SequenceMatcher(None, text, t).ratio() > threshold:
            return True
    return False


def skip_q(text: str) -> bool:
    """Return True for section headers, essay prompts, stray option lines, etc."""
    if len(text) < 10:
        return True
    # Stray option lines like "(A)xxx" or "A.xxx" or "A text"
    if is_option_line(text):
        return True
    # Option continuation lines: Chinese text then (B) or (C) mid-sentence
    if re.search(r"\S\s+[（(][B-D][）)]\S", text):
        return True
    skip_patterns = [
        r"^[一二三四五六七八九十]+[、，,。]",
        r"^[①②③④⑤⑥⑦⑧⑨⑩]",      # circled number continuations
        r"论述题", r"辩论题", r"综合",
        r"^正方[：:]", r"^反方[：:]",
        r"^解释[：:]",
        r"题目类型", r"地球物理题目类型",
        r"^必答题", r"^选答题", r"^创新杯题库",
        r"^\d+\.\d+\s*$",   # "3.1" section headers
        r"^第[一二三四五六七八九十]+",
        r"^[（(][一二三四五六七八九十][）)]",  # （一）sections
    ]
    for pat in skip_patterns:
        if re.search(pat, text):
            return True
    return False


# ─── file parsers ─────────────────────────────────────────────────────────────

def parse_chuangxin_docx(path: str) -> list[dict]:
    """地震学与地震勘探-创新杯.docx
    Section (一): single/multi choice, 20 Qs — format: 01、question（answer）
                 followed by options on next line (tab-separated)
    Section (二): fill-in-blank, 20 Qs
    Section (三)+: essays → skip
    """
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    results = []
    section = None
    i = 0

    while i < len(paras):
        text = paras[i]
        i += 1
        if not text:
            continue

        if re.match(r"[（(][一]", text):
            section = "choice"
            continue
        if re.match(r"[（(][二]", text):
            section = "fill"
            continue
        if re.match(r"[（(][三四五六七八九十]", text) or "论述" in text or "辩论" in text:
            section = "skip"
            continue

        if section == "skip" or section is None:
            continue

        if section == "choice":
            if not re.match(r"^\d{2}[、，,]\s*\S", text):
                continue
            ans_letters = extract_answer_letters(text)
            q_text = clean_q_text(text)
            # Collect tab-separated options from next line
            opts = []
            if i < len(paras):
                next_line = paras[i].strip()
                if next_line and re.match(r"^[A-D][．\.]", next_line):
                    opts = parse_opts_from_single_line(next_line)
                    if not opts:
                        # Try parsing as tab-separated
                        opt_parts = re.split(r"\t+", next_line)
                        for op in opt_parts:
                            m = re.match(r"^([A-D])[．\.\、]\s*(.+)", op.strip())
                            if m:
                                opts.append(f"{m.group(1)}. {m.group(2).strip()}")
                    if opts:
                        i += 1  # consume options line
            if q_text:
                qtype = infer_type(opts, ans_letters or "")
                results.append({"q": q_text, "opts": opts, "a": ans_letters or "", "type": qtype})
            continue

        if section == "fill":
            if not re.match(r"^\d{2}[、，,]\s*\S", text):
                continue
            # Strip numbering from start
            body = re.sub(r"^\d{2}[、，,]\s*", "", text).strip()
            ans, q_text = extract_fill_answer(body)
            if q_text and not skip_q(q_text):
                results.append({"q": q_text, "opts": [], "a": ans, "type": "fill"})

    return results


def parse_chuangxin_qanda(path: str) -> list[dict]:
    """创新杯题库答案&总结.docx — fill and choice questions."""
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    results = []
    i = 0

    while i < len(paras):
        text = paras[i]
        i += 1

        if not text or skip_q(text):
            continue

        # Check for a fill question: has （text_answer）, not just （letter）
        fill_matches = re.findall(r"[（(]([^（）()\n]{2,80})[）)]", text)
        ans_letters = extract_answer_letters(text)

        has_real_fill = any(
            not re.match(r"^[A-D]{1,4}$", a) and len(a) > 1
            for a in fill_matches
        ) if fill_matches else False

        # --- fill question ---
        if has_real_fill and not ans_letters:
            ans, q_text = extract_fill_answer(text)
            q_text = re.sub(r"^\d{1,3}[\.\s。]\s*", "", q_text).strip()
            if q_text and not skip_q(q_text) and len(q_text) > 5:
                results.append({"q": q_text, "opts": [], "a": ans, "type": "fill"})
            continue

        # --- choice question: answer letter in parens ---
        if ans_letters:
            q_text = clean_q_text(text)
            q_text = re.sub(r"^\d{1,3}[\.\s。]\s*", "", q_text).strip()
            if not q_text or skip_q(q_text) or len(q_text) < 5:
                continue
            # Collect option lines that follow
            opts, i = parse_opts_multiline(paras, i)
            qtype = infer_type(opts, ans_letters)
            results.append({"q": q_text, "opts": opts, "a": ans_letters, "type": qtype})
            continue

        # --- numbered question, maybe no inline answer ---
        m_num = re.match(r"^\d{1,3}[\.\s。]\s*(.{6,})", text)
        if m_num:
            body = m_num.group(1).strip()
            if skip_q(body):
                continue
            # Peek ahead for options
            opts, j = parse_opts_multiline(paras, i)
            if opts:
                i = j
                letters = extract_answer_letters(body) or ""
                q_text = clean_q_text(body)
                qtype = infer_type(opts, letters)
                results.append({"q": q_text, "opts": opts, "a": letters, "type": qtype})
            # else: no options, no answer → likely essay intro; skip

    return results


def parse_sheng_sai(path: str) -> list[dict]:
    """6.省赛题（梅、王改版）.docx — 必答题 + 选答题."""
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    results = []
    i = 0

    while i < len(paras):
        text = paras[i]
        i += 1

        if not text or skip_q(text):
            continue

        # Numbered question
        m_num = re.match(r"^\d{1,3}[\.\s。]\s*(.{4,})", text)
        if not m_num:
            # Bare statement with fill blank
            if re.search(r"[（(][^\d（）()\n]{2,60}[）)]", text) and not skip_q(text):
                ans, q_text = extract_fill_answer(text)
                if q_text and not skip_q(q_text):
                    results.append({"q": q_text, "opts": [], "a": ans, "type": "fill"})
            continue

        body = m_num.group(1).strip()
        if skip_q(body):
            continue

        ans_letters = extract_answer_letters(body)

        # Check for fill blanks (non-letter content in parens)
        fill_matches = re.findall(r"[（(]([^（）()\n]{2,60})[）)]", body)
        has_real_fill = any(
            not re.match(r"^[A-D]{1,4}$", a)
            for a in fill_matches
        ) if fill_matches else False

        if has_real_fill and not ans_letters:
            ans, q_text = extract_fill_answer(body)
            q_text = re.sub(r"^\d{1,3}[\.、，,。]\s*", "", q_text).strip()
            if q_text and not skip_q(q_text):
                results.append({"q": q_text, "opts": [], "a": ans, "type": "fill"})
            continue

        # Choice or multi-select
        if ans_letters:
            q_text = clean_q_text(body)
            if not q_text or skip_q(q_text):
                continue
            # Collect option lines
            opts, i = parse_opts_multiline(paras, i)
            qtype = infer_type(opts, ans_letters)
            results.append({"q": q_text, "opts": opts, "a": ans_letters, "type": qtype})
            continue

        # No embedded answer — check for options ahead
        opts, j = parse_opts_multiline(paras, i)
        if opts:
            i = j
            q_text = clean_q_text(body)
            results.append({"q": q_text, "opts": opts, "a": "", "type": "choice"})

    return results


def parse_tmu_带答案(path: str) -> list[dict]:
    """题目（带答案）.docx — numbered fill-in-blank questions."""
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    results = []

    for text in paras:
        if not text or skip_q(text):
            continue
        # Numbered question
        m = re.match(r"^\d{1,3}[\.\s。]\s*(.{6,})", text)
        if not m:
            continue
        body = m.group(1).strip()
        if skip_q(body):
            continue

        ans_letters = extract_answer_letters(body)
        fill_matches = re.findall(r"[（(]([^（）()\n]{2,80})[）)]", body)
        has_real_fill = any(
            not re.match(r"^[A-D]{1,4}$", a) and len(a) > 1
            for a in fill_matches
        ) if fill_matches else False

        if has_real_fill:
            ans, q_text = extract_fill_answer(body)
            q_text = re.sub(r"^\d{1,3}[\.、，,。]\s*", "", q_text).strip()
            if q_text and not skip_q(q_text):
                results.append({"q": q_text, "opts": [], "a": ans, "type": "fill"})
        elif ans_letters and len(ans_letters) == 1:
            # Single-letter embedded — treat as fill (answer is the letter label)
            q_text = clean_q_text(body)
            if q_text and not skip_q(q_text):
                results.append({"q": q_text, "opts": [], "a": ans_letters, "type": "fill"})

    return results


# ─── main ─────────────────────────────────────────────────────────────────────

def main(dry_run: bool = False):
    with open(OUT_JSON, encoding="utf-8") as f:
        existing: dict = json.load(f)

    existing_texts = [q["q"] for q in existing.values()]
    next_id = max(int(k) for k in existing.keys()) + 1

    files_and_parsers = [
        (os.path.join(BASE_DIR, "地震学与地震勘探-创新杯.docx"), parse_chuangxin_docx),
        (os.path.join(BASE_DIR, "创新杯题库答案&总结.docx"),    parse_chuangxin_qanda),
        (os.path.join(BASE_DIR, "6.省赛题（梅、王改版）.docx"), parse_sheng_sai),
        (os.path.join(BASE_DIR, "题目（带答案）.docx"),         parse_tmu_带答案),
    ]

    added = []
    skipped_dup = 0
    skipped_bad = 0

    for path, parser_fn in files_and_parsers:
        print(f"\n📄 {os.path.basename(path)}")
        raw = parser_fn(path)
        print(f"   Raw extracted: {len(raw)}")
        for item in raw:
            q_text = item.get("q", "").strip()
            if not q_text or len(q_text) < 6 or skip_q(q_text):
                skipped_bad += 1
                continue
            if is_duplicate(q_text, existing_texts):
                skipped_dup += 1
                continue

            opts = item.get("opts", [])
            ans = item.get("a", "").strip()
            qtype = item.get("type", "fill")
            # Validate type consistency
            if opts and qtype == "fill":
                qtype = infer_type(opts, ans)
            if not opts and qtype in ("choice", "multi"):
                qtype = "fill"

            category = detect_category(q_text + " " + " ".join(opts))

            new_q = {
                "id": next_id,
                "category": category,
                "type": qtype,
                "section": "",
                "q": q_text,
                "opts": opts,
                "a": ans,
            }
            existing[str(next_id)] = new_q
            existing_texts.append(q_text)
            added.append(new_q)
            print(f"   + [{next_id}] [{category}][{qtype}] {q_text[:60]}")
            next_id += 1

    print(f"\n── Summary ──")
    print(f"  Added:       {len(added)}")
    print(f"  Skipped dup: {skipped_dup}")
    print(f"  Skipped bad: {skipped_bad}")
    print(f"  Total now:   {len(existing)}")

    if dry_run:
        print("\n[DRY RUN] Not writing to disk.")
        return

    with open(OUT_JSON, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)
    print(f"\n✅ Written to {OUT_JSON}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--dry-run", action="store_true", help="Preview without writing")
    args = parser.parse_args()
    main(dry_run=args.dry_run)
